import docx
def n1_1(a, P):

    def ob(p, np, pos):
        b = a[p]
        l = np[p]
        nex = ""
        mi = 99999
        for i in range(len(b)):
            if np[b[i][0]] > l+int(b[i][1:]):
                np[b[i][0]] = l+int(b[i][1:])
            if l+int(b[i][1:]) < mi and b[i][0] not in pos:
                mi = l+int(b[i][1:])
                nex = b[i][0]
        pos+=p
        return nex, np, pos

    def deykstra(x, y, pos = ""):
        np = {x:0}
        nps = x
        for i in a:
            if i not in nps:
                np[i] = 9999
                nps+= i
        n = x
        while n !="":
            b = ob(n, np, pos)
            np = b[1]
            n = b[0]
            pos = b[2]


        return np[y], pos[:pos.find(y)+1]

    k = 0
    p = ""
    for i in range(len(P)):
        b = deykstra(*P[i], p)
        k+=b[0]
        p = b[1]
    print(k)

n1_1({"A": ["B2", "C9", "D4"], "B":["A2",'C3',"E5"], "C":["A9","B2","D6","E10"], "D":["A4","C6","E8"], "E":["B5","C10","D8"]}, [["A", "C"], ["C", "E"]])#нет
def n1_2(a, P):

    def ob(p, np, pos):
        b = a[p]
        l = np[p]
        nex = ""
        mi = 99999
        for i in range(len(b)):
            if np[b[i][0]] > l+int(b[i][1:]):
                np[b[i][0]] = l+int(b[i][1:])
            if l+int(b[i][1:]) < mi and b[i][0] not in pos:
                mi = l+int(b[i][1:])
                nex = b[i][0]
        pos+=p
        return nex, np, pos

    def deykstra(x, y):
        pos = ""
        np = {x:0}
        nps = x
        for i in a:
            if i not in nps:
                np[i] = 9999
                nps+= i
        n = x
        while n !="":
            b = ob(n, np, pos)
            np = b[1]
            n = b[0]
            pos = b[2]

        return np[y]

    k = 0
    for i in range(len(P)):
        k+=deykstra(*P[i])
    print(k)
n1_2({"A": ["B1", "C2", "E4"], "B":["A1",'C4'], "C":["A2","B4","E1"], "D":["E4"], "E":["A4","C1","D4"]}, [["B", "D"]])#верно

def n2():
    for x in range(2):
        for y in range(2):
            for z in range(2):
                for w in range(2):
                    if (not(y <= (x == w)) and (z<=x)) == 1:
                        print(x,y,z,w)
#0111
#0110
#1010
#ответ wxyz

def n4():
    a = "АБВГДЕИН"
    b = {"A":"110", "Б":"01", "И":"000"}
    s = "ВВЕДЕНИЕ"
    def Fano(k):
        m = ["0", "1"]
        l = 1
        while k != l:
            for i in range(len(m)):
                if len(m[i]) == l:
                    m.append(m[i]+"0")
                    m.append(m[i]+"1")
            l+=1
        return m
    def free(a, b):
        rez = []
        c = []
        for i in a:
            for y in range(1, len(a[i])+1):
                c.append(a[i][:y])
        for i in range(len(b)):
            if b[i] not in c:
                t = True
                for z in a:
                    if len(a[z]) < len(b[i]):
                        if a[z] == b[i][:len(a[z])]:
                            t = False
                if t:
                    rez.append(b[i])
        return rez
    kof = 1
    while True:
        m = Fano(kof)
        r = free(b, m)
        if len(r)<(len(a) - len(b)):
            kof+=1
        else:
            m = Fano(kof+1)
            r = free(b, m)
            break

def n5():
    for i in range(100, 1000):
        a = str(i)
        b = int(a[0])*int(a[1])
        c = int(a[1])*int(a[2])
        m = sorted([b, c])
        if str(m[0])+str(m[1]) =="621":
            print(i)
            break
#237
def n6():
    for i in range(1, 10000):
        s = i
        s = s//7
        n = 1
        while s<255:
            if (s+n)%2 == 0:
                s+=11
            n = n+5
        if n == 106:
            print(i)
            break
#175
def n7():
    print(1024*1024*4/8/1024)
#512
def n8():
    k = 0
    for i in "ЛЕВИЙ":
        for x in "ЛЕВИЙ":
            for y in "ЛЕВИЙ":
                for z in "ЛЕВИЙ":
                    for w in "ЛЕВИЙ":
                        if i!= "Й" and (i+x+y+z+w).count("Л") < 2 and (i+x+y+z+w).count("Е") < 2 and (i+x+y+z+w).count("В") < 2 and (i+x+y+z+w).count("И") < 2 and (i+x+y+z+w).count("Й") < 2:
                            if "ЕИ" not in i+x+y+z+w:
                                k+=1
    print(k)
#78

def n10_1():
    k = 0
    a = docx.Document("10.docx")
    par = a.paragraphs
    for i in par:
        k += i.text.count("Мой")
    print(k)
#5
def n10_2():#no
     k = 0
     a = docx.Document("10_demo.docx")
     par = a.paragraphs
     for i in par:
         k += i.text.lower().count("звук")

     print(k)

#мой - 3 прав - 1
def n11():
    a = 8*7//8
    print(a*110)
#770

def n12():
    a = "1"+("9"*100)
    while "19" in a or "299" in a or "3999" in a:
        a = a.replace("19", "2", 1)
        a = a.replace("299", "3", 1)
        a = a.replace("3999", "1", 1)
    print(a)
#39

def n13():
    a = "АБВГ БВД ВГДЕ ГЕИ ДЕЖК ЕИ ИК ЖКЛ КЛМН ЛН МН".split()
    b = {i[0]: i[1:] for i in a}
    def f(x, y):
        if x == "В":
            return 0
        elif x == y:
            return 1
        else:
            k = 0
            s = b[x]
            for i in range(len(s)):
                k+= f(s[i], y)
            return k
    print(f("А", "Н"))
#16

def n14():
    a = 4*(625**9)-25**15+2*(5**11)- 7
    b = ""
    while a>0:
        b+=str(a%5)
        a = a//5
    print(b.count("4"))
#15

def n15():
    ma = 0
    for a in range(100):
        f = True
        for x in range(1000):
            for y in range(1000):
                if not((2*x+3*y < 30) or (x+y >= a)):
                    f = False
        if f:
            ma = a
    print(ma)
#10

def schet():
    def sum(a, b):
      return a[1]+b[1]

    def delenie(m, k = 0):
      if len(m)>1:
        a = delenie(m[:len(m)//2], k)
        b = delenie(m[len(m)//2:], k)
        k+=a[1]
        k+=b[1]
        c = sliyanie(a[0], b[0])
        k+=c[1]
        return c[0], k
      else:
        return m, k


    def sliyanie(a, b):
      k = 0
      m = []
      while len(a)> 0 and len(b)>0:
        if a[0] > b[0]:
          k+= 1
          m.append(b[0])
          b = b[1:]
        else:
          m.append(a[0])
          a = a[1:]
      m+=a
      m+=b
      print(m, k)
      return m, k